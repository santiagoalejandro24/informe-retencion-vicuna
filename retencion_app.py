import streamlit as st
from datetime import datetime
from docx import Document
from io import BytesIO

st.set_page_config(page_title="Informe de Retención - Vicuña", page_icon="📄")
st.title("📄 Generador de Informe de Retención")
st.write("Completá los datos para generar un informe redactado automáticamente.")

hora = st.time_input("⏰ Hora del evento", value=datetime.now().time())
nombre_guardia = st.text_input("👮 Nombre del guardia (Ej: GS Tolaba Fernando)")
nombre_persona = st.text_input("👤 Nombre de la persona retenida")
dni = st.text_input("🪪 DNI o identificación")
empresa = st.text_input("🏢 Empresa (Ej: Transporte Coopera)")
unidad = st.text_input("🔧 Proyecto / Unidad", value="Vicuña Filo del Sol")

tipo_retencion = st.selectbox(
    "📦 Tipo de retención",
    ["Chequeo de bolsos", "Retención vehicular"]
)

motivo = st.text_area("📝 Motivo de la retención", value="retención de exceso de viandas excediendo el límite permitido.")
elementos = st.text_area("📋 Elementos retenidos (uno por línea)", value="""6 jugos Baggios
6 turrones
6 sobres de café
10 sobres de azúcar 
5 saquitos de té
12 tortitas""")

libro = st.text_input("📘 Número de libro de actas", value="1868")
foja = st.text_input("📄 Número de foja", value="103")

if st.button("📤 Generar informe"):
    doc = Document()
    doc.add_paragraph("Señores\nSEGURIDAD PATRIMONIAL\nProyecto Vicuña\nS_/_D\n")

    hora_str = hora.strftime("%H:%M")
    parrafo = (
        f"Se informa que siendo las {hora_str}hs, {nombre_guardia} realiza control de egreso de "
        f"{tipo_retencion.lower()} al Sr. {nombre_persona}, DNI: {dni}, empresa {empresa}, perteneciente a {unidad}. "
        f"Se procede a realizar {motivo.strip()}"
    )
    doc.add_paragraph(parrafo)

    doc.add_paragraph("• " + "\n• ".join([line.strip() for line in elementos.strip().split("\n") if line.strip()]))

    doc.add_paragraph("Realiza descargo en acta correspondiente. Se procede al resguardo de los mismos para su posterior decomiso.")
    doc.add_paragraph(f"Se deja constancia en libro de actas N.º {libro} foja {foja}.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    nombre_archivo = f"Informe_retención_{nombre_persona.replace(' ', '_')}_{datetime.now().strftime('%d-%m-%Y')}.docx"

    st.success("✅ Informe generado correctamente.")
    st.download_button(
        label="📥 Descargar informe redactado",
        data=buffer,
        file_name=nombre_archivo,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
